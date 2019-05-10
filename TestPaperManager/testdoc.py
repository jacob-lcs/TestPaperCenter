from docx import Document
from docx.shared import Inches


def make_new_file():
    document = Document()
    document.add_heading('试题导出', 0)  # 好像是标题一样的东西哦
    p = document.add_paragraph('这 是一个p')
    p.add_run('bold').bold = True  # 加粗
    p.add_run(' and some ')  # 正常文字
    p.add_run('italic.').italic = True  # 斜体

    document.add_heading('Heading, level 1', level=1)  # 一级标题样式
    document.add_paragraph('Intense quote', style='Intense Quote')  # 明显引用

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )  # 项目符号列表
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )  # 数字编号

    document.add_picture('achievement_gurenl.png', width=Inches(1.25))  # 插入图片！有用的功能！

    # ----------插入表格儿------------------------
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )  # 插入表格儿

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    # -------------------------------------------

    document.add_page_break()  # 加入分页符

    document.save('demo.docx')


def read_a_file():
    import zipfile
    f = zipfile.ZipFile('demo.docx', 'r')
    # 解压文件
    for i in f.namelist():
        print(i)
        print(f.extract(i))


if __name__ == '__main__':
    read_a_file()
